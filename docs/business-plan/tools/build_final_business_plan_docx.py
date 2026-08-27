from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


FONT_NAME = "휴먼명조"
TABLE_WIDTH_DXA = 10080
TABLE_INDENT_DXA = 120
BLACK = "000000"
WHITE = "FFFFFF"
GRAY = "666666"


def set_run_font(run, size=11, bold=False, color=BLACK, italic=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)


def set_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=6,
    line_spacing=1.6,
    keep_with_next=False,
    keep_together=False,
):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, color=WHITE):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, *, size=8, color=BLACK):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name in ("top", "bottom", "start", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge_name}")
        edge = borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def remove_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name in ("top", "bottom", "start", "end", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def apply_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for tr in table._tbl.tr_lst:
        col_index = 0
        for tc in tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            span_node = tc_pr.find(qn("w:gridSpan"))
            span = int(span_node.get(qn("w:val"))) if span_node is not None else 1
            width = sum(widths[col_index : col_index + span])
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            col_index += span


def set_cell_text(
    cell,
    text,
    *,
    size=9.5,
    bold=False,
    color=BLACK,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing=1.35,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(
        paragraph,
        alignment=alignment,
        before=0,
        after=0,
        line_spacing=line_spacing,
        keep_together=True,
    )
    lines = str(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_name(document, name, *, size, bold, before, after, line_spacing, alignment, color=BLACK):
    style = document.styles[name]
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    for attr in ("ascii", "hAnsi", "cs"):
        style._element.rPr.rFonts.set(qn(f"w:{attr}"), FONT_NAME)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.alignment = alignment
    style.paragraph_format.keep_with_next = True


def configure_document(document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    style_name(
        document,
        "Normal",
        size=11,
        bold=False,
        before=0,
        after=6,
        line_spacing=1.6,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    style_name(
        document,
        "Heading 1",
        size=16,
        bold=True,
        before=0,
        after=8,
        line_spacing=1.2,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    style_name(
        document,
        "Heading 2",
        size=13,
        bold=True,
        before=8,
        after=5,
        line_spacing=1.2,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    style_name(
        document,
        "Heading 3",
        size=11.5,
        bold=True,
        before=6,
        after=3,
        line_spacing=1.2,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    style_name(
        document,
        "Caption",
        size=9,
        bold=False,
        before=3,
        after=6,
        line_spacing=1.2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    for style_name_value in ("List Bullet", "List Number"):
        style = document.styles[style_name_value]
        style.font.name = FONT_NAME
        style.font.size = Pt(10.5)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
        style.paragraph_format.left_indent = Mm(8.5)
        style.paragraph_format.first_line_indent = Mm(-4.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.45

    set_header_footer(section)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        before=0,
        after=0,
        line_spacing=1.0,
    )
    run = p.add_run("제13회 전국 ICT융합 공모전  |  비자부기")
    set_run_font(run, size=8.5, color=GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=0,
        line_spacing=1.0,
    )
    run = p.add_run("비자부기 사업계획서  |  ")
    set_run_font(run, size=8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fld_run = OxmlElement("w:r")
    fld_rpr = OxmlElement("w:rPr")
    fld_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fld_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    fld_rpr.append(fld_fonts)
    fld_run.append(fld_rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    fld_run.append(text)
    fld.append(fld_run)
    p._p.append(fld)

    # 공식 양식의 표지는 머리말·꼬리말 없이 사용한다.
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""


def add_body(document, text, *, bold_lead=None, after=6, size=11, italic=False):
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, after=after, line_spacing=1.6)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=size, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, size=size, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, italic=italic)
    return paragraph


def add_bullet(document, text, *, size=10.5):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    return paragraph


def add_numbered(document, text, *, size=10.5):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_run_font(run, size=9)
    return paragraph


def add_standard_table(
    document,
    headers,
    rows,
    widths,
    *,
    font_size=9.2,
    first_col_bold=False,
    center_columns=None,
):
    if center_columns is None:
        center_columns = set()
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_fill(cell, BLACK)
        set_cell_border(cell, size=8)
        set_cell_margins(cell, top=100, start=100, bottom=100, end=100)
        set_cell_text(
            cell,
            header,
            size=font_size,
            bold=True,
            color=WHITE,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.25,
        )
    set_repeat_table_header(table.rows[0])
    for r, values in enumerate(rows, start=1):
        for c, value in enumerate(values):
            cell = table.cell(r, c)
            set_cell_fill(cell, WHITE)
            set_cell_border(cell, size=8)
            set_cell_margins(cell, top=100, start=110, bottom=100, end=110)
            alignment = WD_ALIGN_PARAGRAPH.CENTER if c in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(
                cell,
                value,
                size=font_size,
                bold=first_col_bold and c == 0,
                alignment=alignment,
                line_spacing=1.35,
            )
    for row in table.rows:
        prevent_row_split(row)
    apply_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metric_strip(document, items):
    count = len(items)
    widths = [TABLE_WIDTH_DXA // count] * count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = document.add_table(rows=2, cols=count)
    set_repeat_table_header(table.rows[0])
    for index, (value, label) in enumerate(items):
        top = table.cell(0, index)
        bottom = table.cell(1, index)
        for cell in (top, bottom):
            set_cell_fill(cell, WHITE)
            set_cell_border(cell, size=10)
            set_cell_margins(cell, top=90, start=80, bottom=90, end=80)
        set_cell_text(
            top,
            value,
            size=14,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.1,
        )
        set_cell_text(
            bottom,
            label,
            size=8.8,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.25,
        )
    for row in table.rows:
        prevent_row_split(row)
    apply_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_compact_flow(document, steps, footer=None):
    count = len(steps)
    if count == 5:
        box_width, arrow_width = 1776, 300
    elif count == 4:
        box_width, arrow_width = 2250, 360
    elif count == 3:
        box_width, arrow_width = 3060, 450
    else:
        raise ValueError("compact flow supports 3 to 5 steps")
    widths = []
    for index in range(count):
        widths.append(box_width)
        if index < count - 1:
            widths.append(arrow_width)
    assert sum(widths) == TABLE_WIDTH_DXA

    table = document.add_table(rows=2, cols=len(widths))
    set_repeat_table_header(table.rows[0])
    for step_index, (heading, body) in enumerate(steps):
        col = step_index * 2
        header_cell = table.cell(0, col)
        body_cell = table.cell(1, col)
        set_cell_fill(header_cell, BLACK)
        set_cell_fill(body_cell, WHITE)
        set_cell_border(header_cell, size=10)
        set_cell_border(body_cell, size=10)
        set_cell_margins(header_cell, top=100, start=70, bottom=100, end=70)
        set_cell_margins(body_cell, top=130, start=70, bottom=130, end=70)
        set_cell_text(
            header_cell,
            heading,
            size=9.2,
            bold=True,
            color=WHITE,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.25,
        )
        set_cell_text(
            body_cell,
            body,
            size=8.8,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.3,
        )
        if step_index < count - 1:
            arrow_cell = table.cell(0, col + 1).merge(table.cell(1, col + 1))
            set_cell_fill(arrow_cell, WHITE)
            remove_cell_border(arrow_cell)
            set_cell_margins(arrow_cell, top=0, start=0, bottom=0, end=0)
            set_cell_text(
                arrow_cell,
                "→",
                size=15,
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing=1.0,
            )
    for row in table.rows:
        prevent_row_split(row)
    apply_table_geometry(table, widths)

    if footer:
        p = document.add_paragraph()
        set_paragraph_format(
            p,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            before=4,
            after=4,
            line_spacing=1.25,
        )
        run = p.add_run(footer)
        set_run_font(run, size=9, bold=True)
    return table


def add_picture(document, path, *, width_mm, alt_text, caption):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=2,
        after=0,
        line_spacing=1.0,
        keep_with_next=True,
    )
    run = p.add_run()
    shape = run.add_picture(str(path), width=Mm(width_mm))
    shape._inline.docPr.set("descr", alt_text)
    add_caption(document, caption)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLACK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.extend([rfonts, color, underline, size])
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_source(document, number, title, organization, date_text, url):
    paragraph = document.add_paragraph()
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        before=0,
        after=4,
        line_spacing=1.3,
        keep_together=True,
    )
    lead = paragraph.add_run(f"[{number}] {organization}, ")
    set_run_font(lead, size=9, bold=True)
    title_run = paragraph.add_run(f"「{title}」, {date_text}. ")
    set_run_font(title_run, size=9)
    add_hyperlink(paragraph, "원문 보기", url)


def add_cover(document):
    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=16,
        after=8,
        line_spacing=1.0,
    )
    run = p.add_run("- 2026년 지역주도 디지털혁신지원사업 -")
    set_run_font(run, size=12, bold=True)

    p = document.add_paragraph()
    set_paragraph_format(
        p,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=10,
        after=0,
        line_spacing=1.35,
    )
    run = p.add_run("제13회 전국 ICT융합 공모전\n사업계획서")
    set_run_font(run, size=24, bold=True)

    spacer = document.add_paragraph()
    set_paragraph_format(spacer, before=0, after=0, line_spacing=1.0)
    spacer.paragraph_format.space_before = Pt(105)

    table = document.add_table(rows=3, cols=2)
    labels = ["참   가   분   야", "작품명(사업명)", "신 청 인 ( 팀 장 )"]
    values = ["디지털 시제품", "비자부기(visa-bugi)", "김태은"]
    for row_index, (label, value) in enumerate(zip(labels, values)):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)
        for cell in (label_cell, value_cell):
            remove_cell_border(cell)
            set_cell_fill(cell, WHITE)
            set_cell_margins(cell, top=130, start=90, bottom=130, end=90)
        set_cell_text(
            label_cell,
            label,
            size=12,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            line_spacing=1.2,
        )
        set_cell_text(
            value_cell,
            f":  {value}",
            size=12,
            bold=row_index == 1,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=1.2,
        )
    apply_table_geometry(table, [3550, 6530], indent=0)


def add_picture_to_cell(cell, path, *, width_mm, alt_text, caption):
    paragraph = cell.add_paragraph()
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=1,
        after=1,
        line_spacing=1.0,
    )
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Mm(width_mm))
    shape._inline.docPr.set("descr", alt_text)
    caption_paragraph = cell.add_paragraph()
    set_paragraph_format(
        caption_paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        before=1,
        after=0,
        line_spacing=1.0,
    )
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=8.2)


def add_official_summary(document, assets):
    document.add_heading("1. 요약", level=1)
    rows = [
        (
            "제품\n(아이디어)명",
            "비자부기(visa-bugi) — 외국인 주민의 비자 요건·서류·일정·지원기관을 한 번에 관리하는 충북형 AI 비자 동행 서비스",
        ),
        (
            "제품\n(아이디어)\n소개",
            "충북 외국인 주민의 다국어 질문, OCR 신청서 상태, 공식 공고 일정을 분석하여 비자정보·보완서류·행정일정·전문기관 중 적절한 다음 행동으로 연결하는 모바일 우선 웹 서비스이다. AI 챗봇은 질문의 의도·비자·지역·위험 신호를 파악하고, OCR은 서류를 완성·확인 필요·누락·수동 확인으로 구분한다. 자격·점수와 날짜는 검수된 데이터와 결정론적 규칙으로 계산한다.",
        ),
        (
            "제품\n(아이디어)\n사업성",
            "개인에게 핵심 기능을 무료로 제공하고, 대학·기업·지자체·외국인지원기관의 사용료로 AI 처리비와 행정데이터 갱신비를 충당하는 B2C 기반 B2G/B2B 모델이다. OCR 분석비는 현재 시험 가정에서 문서 1건당 약 0.009달러로 추산되어 낮은 처리비로 외국인 주민의 비자 준비와 충북 정착을 지원할 수 있다. [13]",
        ),
        (
            "제품\n(아이디어)\n가치",
            "질문→근거 조회→OCR 서류 점검→공고 일정 연결→기관 안내를 하나의 사용자 여정으로 통합한다. 구축 대상 공식 원천문서 24건 전체를 파싱해 원문 계층·근거 계층·규칙 계층·서비스 계층을 구성했으며, 684개 원천→서비스 매핑으로 결과를 원문까지 추적한다.",
        ),
        (
            "제품\n(아이디어)\n기대효과",
            "정보 탐색시간과 필수서류 누락을 줄이고 외국인 인재의 취업·비자 전환·지역 정착을 연결한다. 대학·기업·지원기관에는 반복 상담 부담 완화와 동일한 최신 근거를 공유하는 업무 기반을 제공한다.",
        ),
    ]

    table = document.add_table(rows=6, cols=2)
    for row_index, (label, value) in enumerate(rows):
        left = table.cell(row_index, 0)
        right = table.cell(row_index, 1)
        for cell in (left, right):
            set_cell_fill(cell, WHITE)
            set_cell_border(cell, size=9)
            set_cell_margins(cell, top=70, start=100, bottom=70, end=100)
        set_cell_text(
            left,
            label,
            size=9.0,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.15,
        )
        set_cell_text(
            right,
            value,
            size=8.6,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=1.2,
        )

    image_left = table.cell(5, 0)
    image_right = table.cell(5, 1)
    for cell in (image_left, image_right):
        set_cell_fill(cell, WHITE)
        set_cell_border(cell, size=9)
        set_cell_margins(cell, top=55, start=90, bottom=55, end=90)
    set_cell_text(
        image_left,
        "관련 이미지",
        size=9.0,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing=1.15,
    )
    image_right.text = ""
    add_picture_to_cell(
        image_right,
        assets / "01-dashboard-mockup.png",
        width_mm=88,
        alt_text="준비 현황과 비자 여정, 다음 할 일을 보여주는 비자부기 대시보드",
        caption="그림 1. 비자부기 대시보드 화면",
    )
    for row in table.rows:
        prevent_row_split(row)
    apply_table_geometry(table, [1850, 8230])


def add_page_break(document):
    document.add_page_break()


def build(output_path):
    document = Document()
    configure_document(document)
    assets = Path(__file__).resolve().parents[1] / "assets"

    # 1쪽 — 공식 양식 표지
    add_cover(document)

    # 2쪽 — 공식 양식 1. 요약
    add_page_break(document)
    add_official_summary(document, assets)

    # 3쪽 — 기획 배경: 문제와 정책
    add_page_break(document)
    document.add_heading("2. 기획 배경", level=1)
    document.add_heading("2-1. 문제 정의", level=2)
    add_body(
        document,
        "외국인 주민이 비자를 준비하려면 먼저 자신의 상황에 맞는 체류자격을 찾아야 한다. 신청 단계에서는 서류 작성 주체와 보완항목을 구분해야 하고, 공고 접수기간과 개인별 행정기한도 따로 관리해야 한다. 어느 하나라도 놓치면 신청이 늦어질 수 있다.",
    )
    add_body(
        document,
        "공식 자료는 PDF·HWPX·HWP의 본문과 표, 각주, 붙임서식에 흩어져 있다. 검색이나 상담으로 답을 얻더라도 필요한 서류와 마감일은 사용자가 다시 정리해야 한다. 비자부기는 AI 챗봇, OCR, 공고 스케줄링을 연결해 질문에 대한 답을 체크리스트와 일정으로 이어준다.",
    )
    add_metric_strip(
        document,
        [
            ("86,490명", "2025년 충북 장기체류외국인"),
            ("27,705명", "충북 E계열 취업 외국인"),
            ("14,303명", "2026년 4월 충북 외국인 유학생"),
            ("13개사", "F-2-R 연계 면접 기업"),
        ],
    )
    add_caption(document, "표 1. 충북 외국인 주민의 취업·유학·정착 수요 [3][8][9]")
    document.add_heading("2-2. 정책·시장 배경", level=2)
    add_body(
        document,
        "법무부 통계에 따르면 2025년 말 국내 체류외국인은 전년 대비 5.0% 증가했으며, 국내 외국인 유학생도 17.1% 늘었다. 주요 국적은 중국·베트남·미국·태국·우즈베키스탄 등으로 다양해 외국인 주민이 이해하기 쉬운 다국어 행정 안내의 필요성이 커지고 있다. [1]",
    )
    add_body(
        document,
        "2025년 충북의 장기체류외국인은 86,490명이다. 이 가운데 E계열 체류자격으로 취업한 외국인은 27,705명이며, 2026년 4월 기준 외국인 유학생은 14,303명에 이른다. [8][9]",
        after=3,
    )
    add_body(
        document,
        "충북은 F-2-R과 E-7-4R 등 지역특화형 비자 사업을 계속 추진하고 있다. 2026년 외국인 유학생 채용박람회에는 유학생 약 700명과 도내 기업 32개사가 참여했다. 인구감소지역에 있는 13개 기업은 F-2-R 전환을 희망하는 유학생과 현장 면접을 진행했다. [2][3]",
        after=3,
    )
    add_body(
        document,
        "채용 이후에는 비자 전환과 지역 정착이 이어진다. 이 과정에서 외국인 인재가 이탈하지 않도록 돕는 일은 충북의 산업인력 확보와 직결되는 디지털 행정 과제다.",
    )

    # 4쪽 — 제안 동기
    add_page_break(document)
    document.add_heading("2-3. 제안 동기", level=2)
    add_body(
        document,
        "비자 안내의 정확성은 화면보다 원천 데이터에서 결정된다. 팀은 서비스 개발에 앞서 공식 공고·지침·서식 24건을 수집했다. PDF·HWPX·HWP에 포함된 본문과 표, 각주, 붙임서식의 구조를 분석하고 추출 결과를 사람이 원문과 다시 대조했다.",
    )
    add_body(
        document,
        "검수를 마친 자료는 자격요건, 점수표, 절차, 제출서류, 쿼터, 변경이력으로 구분했다. 이후 원문, 근거, 규칙, 서비스의 4계층으로 정리해 13개 관계형 테이블에 저장했다. 현재 684개의 근거 연결을 통해 화면에 표시된 정보가 어느 문서에서 나왔는지 확인할 수 있다.",
    )
    add_body(
        document,
        "이 과정에서 일반적인 텍스트 추출만으로는 행정문서를 정확하게 처리하기 어렵다는 사실을 확인했다. 표의 병합 셀과 각주, 구버전 HWP, 문서 간 페이지 차이뿐 아니라 중첩된 AND/OR 조건과 최소점수, 함께 적용할 수 없는 가점, 적용기간까지 해석해야 했다. 문서에 값이 없는 ‘미기재’와 아직 확인하지 못한 ‘미확인’도 구분했다.",
    )
    add_body(
        document,
        "실제 E-7-4R 공고를 검수할 때는 벌금 감점 기준이 본문·점수표와 붙임서식에서 서로 다르게 기재된 사례를 발견했다. 일부 HWPX 문서에서는 빈 개체 때문에 PDF와 원문 페이지가 어긋났다. 팀은 이러한 값을 임의로 통합하지 않았다. 서로 다른 근거를 함께 보존하고 자동 계산에서 제외해 담당자가 다시 확인하도록 했다.",
    )
    document.add_heading("2-4. 기존 방식의 빈칸", level=2)
    add_standard_table(
        document,
        ["기존 서비스·방식", "주요 역할", "비자부기가 보완하는 지점"],
        [
            ["하이코리아 [11]", "체류 관련 전자민원·신청·조회", "신청 전 개인별 요건·서류·일정 준비를 추적"],
            ["정부24 [12]", "범정부 민원 신청과 공식 절차 안내", "비자별 공고·서식·지원기관을 하나의 여정으로 연결"],
            ["Study in Chungbuk [10]", "유학생 중심 다국어 정보와 AI 번역 챗봇", "근로자·지역특화형 비자까지 규칙 계산과 행동관리로 확장"],
            ["범용 생성형 AI", "질문 응답과 쉬운 설명", "자격 계산은 검수된 규칙으로 분리하고 근거·현행성을 표시"],
            ["비자부기", "요건 확인부터 준비·추적까지 연결", "제출·최종 판정은 공식기관이 수행하도록 책임을 구분"],
        ],
        [1800, 3500, 4780],
        font_size=8.8,
        first_col_bold=True,
    )
    add_body(
        document,
        "중복성 측면에서도 기존 공식 서비스의 민원 접수·정보 제공을 복제하지 않는다. 비자부기는 신청 전 준비·추적, OCR 서류 점검, 공고 일정 연결, 위험상황 라우팅이라는 빈칸을 보완한다.",
        after=0,
    )

    # 5쪽 — 목표 사용자와 전체 흐름
    add_page_break(document)
    document.add_heading("3. 기획 세부설명", level=1)
    document.add_heading("3-1. 목표 사용자와 대표 시나리오", level=2)
    add_body(
        document,
        "비자부기의 핵심 사용자는 충북에서 취업과 정착을 준비하는 외국인 유학생, 그리고 E-7-4R 등으로 체류자격 전환을 준비하는 외국인 근로자다. 이들을 채용하는 충북 기업의 인사 담당자와 대학 국제교류부서·가족센터·외국인지원센터·지자체의 상담 담당자도 같은 정보를 확인하는 기관 사용자로 본다.",
    )
    add_body(
        document,
        "대표 시나리오는 충북 소재 기업에 취업한 유학생이 F-2-R 전환을 검토하거나, E-9 근로자가 E-7-4R 전환을 준비하는 과정이다. 비자부기는 질문 이해, 서류 점검, 공고 일정 연결, 위험상황 기관 연결을 하나의 AI 사용자 여정으로 잇는다. 사용자는 공고문이나 신청서를 촬영해 올릴 수도 있다.",
    )
    add_body(
        document,
        "AI는 사용자의 질문에서 경로를 고르고, 자격·점수·연락처·행정일정은 검수된 데이터와 코드가 계산한다. 조건이 불명확하면 통과시키지 않고 ‘검토 필요’로 남기고, 기준일 근거가 없는 상대 일정은 임의로 날짜를 만들지 않는다. 위험 신호가 나오면 일반 답변을 멈추고 전문기관으로 연결하며, 로그인 전 입력값은 현재 브라우저 세션에서만 쓴다.",
    )
    add_caption(document, "그림 2. 비자부기 전체 서비스 흐름")
    add_compact_flow(
        document,
        [
            ("1. 상황 입력", "다국어 질문\nOCR·공식 공고"),
            ("2. AI 스크리닝", "위험·비자·지역\n현재 단계 분류"),
            ("3. 경로 선택", "비자·서류·일정\n위험기관 라우팅"),
            ("4. 근거 검증", "출처·적용기간\n연락처·날짜 확인"),
            ("5. 행동 연결", "체크리스트·캘린더\n기관 연락처"),
        ],
    )

    # 6쪽 — 서비스 기능
    add_page_break(document)
    document.add_heading("3-2. 서비스 구성과 핵심 기능", level=2)
    add_standard_table(
        document,
        ["기능", "사용자에게 제공하는 결과", "판정·저장 원칙"],
        [
            ["AI 챗봇", "의도·비자·지역·위험 신호를 파악해 필요한 경로로 연결", "제한형 데이터 조회만 허용"],
            ["위험 라우팅", "임금체불·산재·폭행·불법취업 등 전문기관 우선 안내", "기관명·연락처를 원문 그대로 제공"],
            ["OCR 서류 점검", "완성·확인 필요·누락·수동 확인과 작성 주체를 구분", "실제 개인정보 대신 필드 상태만 활용"],
            ["공고 스케줄링", "접수기간·절차·개인 기준일을 캘린더 일정으로 등록", "근거 없는 날짜를 생성하지 않음"],
            ["요건·점수 계산", "필수조건·대체조건·점수·부족 항목을 구분", "검수된 AND/OR 규칙으로 계산"],
            ["서류·기관 연결", "체크리스트와 충북 시군별 전화·길찾기 제공", "GPS 좌표·OCR 원본은 비저장"],
        ],
        [2200, 4830, 3050],
        font_size=8.8,
        first_col_bold=True,
    )
    document.add_heading("사용자에게 보이는 결과", level=3)
    add_body(
        document,
        "홈 화면은 준비 현황, 현재 단계, 다음 할 일을 한눈에 보여준다. 각 항목은 출처 문서·페이지·적용기간·마지막 검증일과 연결되어 사용자가 근거를 다시 확인할 수 있다.",
    )
    add_body(
        document,
        "즉, 비자부기는 질문을 알아듣고 서류를 읽으며 행정기한을 놓치지 않도록 돕는 충북형 AI 비자 동행 서비스다.",
    )

    # 7쪽 — 데이터와 판정
    add_page_break(document)
    document.add_heading("3-3. AI 행정·안전 라우팅과 데이터 신뢰 구조", level=2)
    add_body(
        document,
        "AI는 사용자의 질문에서 의도와 언어, 지역, 비자 유형, 위험 신호를 파악한다. 답변에 포함할 요건과 절차, 제출서류, 기관은 검수된 데이터로 제한한다. 자격과 점수, 날짜 계산은 코드가 담당한다.",
        after=3,
    )
    add_body(
        document,
        "이를 위해 공식 원천문서 24건을 원문, 근거, 규칙, 서비스의 4계층으로 구조화했다. 질문이 들어오면 정해진 유형으로 분류한 뒤 해당 계층의 데이터만 조회한다.",
    )
    add_caption(document, "그림 3. 공식 문서에서 서비스까지 이어지는 4계층 데이터 구조")
    add_standard_table(
        document,
        ["계층", "관리 대상", "검증 방식", "다음 계층과의 관계"],
        [
            ["1. 원문 계층", "공식 공고·지침·붙임서식 24건", "본문·표·각주·페이지 보존", "원문 구간을 근거 레코드에 연결"],
            ["2. 근거 계층", "원문 구절·표 셀·적용기간·검증일", "684개 원천→서비스 매핑", "각 근거를 요건·서류·일정 규칙에 연결"],
            ["3. 규칙 계층", "AND/OR 요건·점수·서류·절차·일정", "사람 검수와 무결성 검사", "검증을 통과한 규칙만 서비스에서 사용"],
            ["4. 서비스 계층", "챗봇·OCR·캘린더·기관 안내", "출처·날짜·연락처 재검증", "체크리스트와 일정, 기관 연락처로 표시"],
        ],
        [1650, 2970, 2670, 2790],
        font_size=8.5,
        first_col_bold=True,
    )
    document.add_heading("기능별 역할 구분", level=3)
    add_bullet(document, "비자 챗봇은 의도·위험·언어를 분류하고 쉬운 표현으로 설명한다. 요건·절차·서류·기관 조회 범위는 검수된 데이터로 제한한다.", size=9.2)
    add_bullet(document, "OCR 도우미는 선택한 필드의 의미를 설명한다. 서식 템플릿, 작성 주체, 누락 상태는 정해진 규칙으로 판정한다.", size=9.2)
    add_bullet(document, "공고 일정은 사용자 상황에 맞춰 안내한다. 날짜 계산과 변경 비교에는 확정일, 기준일, 공식 offset만 사용한다.", size=9.2)
    add_body(
        document,
        "이 구조는 13개 관계형 테이블과 684개 원천→서비스 매핑으로 답변을 원문까지 추적한다. 데이터에 없는 답변이나 연락처는 생성하지 않고, 자격 AND/OR·점수와 날짜는 코드로 계산하며 검토 항목은 자동 통과시키지 않는다.",
    )

    # 8쪽 — 통합 기능과 책임 경계
    add_page_break(document)
    document.add_heading("3-4. AI 챗봇·OCR·공고 일정의 통합", level=2)
    add_body(
        document,
        "AI 챗봇은 매 질문에서 위험 신호·사용자 유형·충북 시군·관심 비자를 먼저 파악한다. 위험 신호가 있으면 일반 답변을 중단하고 검수된 라우팅 규칙으로 전문기관을 안내하며, 일반 질문은 요건·절차·서류·쿼터·기관 조회 도구로 연결한다.",
    )
    add_body(
        document,
        "OCR은 신청서의 항목을 완성, 확인 필요, 누락, 수동 확인으로 나눈다. 신청자와 고용주, 기관 중 누가 작성해야 하는지도 함께 표시한다.",
        after=3,
    )
    add_body(
        document,
        "공고 스케줄러는 접수 시작일과 마감일, 절차, 적용기간을 목표 비자에 연결한다. 확정일은 그대로 표시한다. 개인 기준일이 필요한 일정은 기준일을 확인한 뒤 계산하고, 공고가 바뀌면 이전 일정과 달라진 근거를 함께 보여준다.",
    )
    add_caption(document, "그림 4. 비자 여정 단계별 서비스 책임 경계")
    add_standard_table(
        document,
        ["단계", "내용", "비자부기 역할", "최종 책임 주체"],
        [
            ["1. 추적", "요건·진행 현황 확인", "규칙 계산으로 자가진단 지원", "비자부기(참고용)"],
            ["2. 준비", "서류·OCR·일정 정리", "문서·OCR·초안 지원", "비자부기 + 사용자"],
            ["3. 제출", "공식 링크·제출처 안내", "안내까지만", "공식 접수기관"],
            ["4. 판정", "승인·불허·보완", "관여하지 않음", "관할 행정기관"],
        ],
        [1300, 2700, 3300, 2780],
        font_size=9.1,
        first_col_bold=True,
    )
    document.add_heading("개인정보와 안전 원칙", level=3)
    add_bullet(document, "GPS 사용은 선택이며 좌표를 저장하지 않고, 거부해도 시·군을 직접 선택할 수 있다.")
    add_bullet(document, "OCR 원본과 여권번호·주소 등 민감정보는 최소 처리·비저장을 기본값으로 둔다.")
    add_bullet(document, "위험상황은 일반 대화에서 분리해 전문기관 또는 전국 단위 긴급기관으로 라우팅한다.")
    add_bullet(document, "모든 자가진단 화면에 공식 결정 대체가 아니라는 책임 경계를 표시한다.")

    # 9쪽 — 대표 사용자 여정 2종
    add_page_break(document)
    document.add_heading("3-5. 대표 사용자 여정", level=2)
    document.add_heading("① 이주노동자: 비자 준비와 노동권 보호", level=3)
    add_body(
        document,
        "응우옌 반 A · 27세 · 베트남 · E-9 · 음성군 제조업체 근무 2년 차. E-7-4R 전환 가능성과 임금체불 문제를 함께 확인하는 상황이다.",
        size=9.5,
        after=3,
    )
    add_caption(document, "그림 5. 이주노동자의 E-7-4R 준비와 위험상황 연결")
    add_compact_flow(
        document,
        [
            ("1. AI 질문 이해", "비자·지역·현재단계\n위험 표현 분류"),
            ("2. 위험 우선", "임금체불 감지\n일반 답변 중단"),
            ("3. 기관 라우팅", "사용자·지역 대조\n연락처 원문 안내"),
            ("4. 비자 준비", "요건·OCR 서류\n부족 항목 확인"),
            ("5. 일정 연결", "공고 마감·방문일\n캘린더 관리"),
        ],
    )
    add_bullet(document, "AI가 비자 문의와 임금체불 신호를 함께 파악하면 일반 답변보다 지역·사용자 유형에 맞는 노동 전문기관 연결을 우선한다.", size=9.2)
    add_bullet(document, "이후 검수된 E-7-4R 요건·OCR 서류 상태·공고 마감일을 체크리스트와 캘린더 일정에 반영한다.", size=9.2)

    document.add_heading("② 외국인 유학생: 유학→취업→정착", level=3)
    add_body(
        document,
        "바트 체첵 · 22세 · 몽골 · D-2 · 충북 소재 대학 재학. 시간제취업 허가와 졸업 후 충북 정착 경로를 함께 준비하는 상황이다.",
        size=9.5,
        after=3,
    )
    add_caption(document, "그림 6. 유학생의 유학→취업→정착 준비 과정")
    add_compact_flow(
        document,
        [
            ("1. AI 질문 이해", "학적·비자·지역\n문의 의도 분류"),
            ("2. OCR 서류 점검", "작성 주체·누락\n수동 확인 구분"),
            ("3. 필드별 설명", "누가·왜 작성하는지\n6개 언어 안내"),
            ("4. 공고 일정 연결", "접수기간·제출기한\n출처와 함께 표시"),
            ("5. 행동 관리", "체크리스트·캘린더\n담당기관 연결"),
        ],
    )
    add_bullet(document, "OCR이 작성 주체·필수 여부·누락·수동 확인을 구분하고, 챗봇은 실제 개인정보 대신 필드 상태로 작성 방법을 설명한다.", size=9.2)
    add_bullet(document, "목표 비자의 공고·절차·기관을 조회해 누락서류는 체크리스트에, 확인된 마감일은 캘린더에 연결한다.", size=9.2)

    # 10쪽 — 차별점과 데이터 자산
    add_page_break(document)
    document.add_heading("3-6. 차별점과 기술적 우수성", level=2)
    add_standard_table(
        document,
        ["비교 항목", "하이코리아·정부24", "Study in Chungbuk", "비자부기"],
        [
            ["핵심 역할", "공식 민원 신청·조회", "유학생 정보·다국어 안내", "AI 기반 신청 전 준비·추적"],
            ["질문 처리", "절차를 사용자가 탐색", "정보·번역 중심", "의도·위험·비자·지역 라우팅"],
            ["서류 지원", "서식 제공", "콘텐츠 안내", "OCR 필드 상태·작성 주체 설명"],
            ["일정 지원", "공식 절차별 확인", "정보 탐색 중심", "공고·기준일을 캘린더 일정과 할 일로 등록"],
            ["안전 개입", "공식 접수기관 안내", "일반 문의 안내", "위험 신호 감지 후 전문기관 우선 연결"],
            ["결과 형태", "민원 신청·조회", "정보·링크", "체크리스트·캘린더·기관 연락처"],
        ],
        [1900, 2440, 2440, 3300],
        font_size=8.4,
        first_col_bold=True,
    )
    add_caption(document, "표 2. 공식 서비스와 비자부기의 역할 비교 [10][11][12]")
    document.add_heading("구축된 데이터·검증 자산", level=3)
    add_metric_strip(
        document,
        [
            ("13개", "관계형 테이블"),
            ("1,101행", "공통 데이터 레코드"),
            ("24건", "원천 문서"),
            ("684건", "원천→공통 매핑"),
        ],
    )
    add_metric_strip(
        document,
        [
            ("111건", "자격조건"),
            ("77건", "제출서류"),
            ("97건", "기관 연락처"),
            ("407개", "자동 테스트 통과"),
        ],
    )
    add_caption(document, "표 3. 2026년 8월 26일 팀 데이터 저장소 검증 스냅샷 [6]")
    add_body(
        document,
        "질문과 OCR 서류 상태, 공고를 함께 해석하는 AI 라우팅 구조가 비자부기의 핵심 차별점이다. 위험 대응, 비자정보, 서류 점검, 일정 관리 중 필요한 경로를 고르고 결과를 체크리스트·캘린더·기관 연락처로 연결한다.",
    )
    add_body(
        document,
        "공식 원천문서 24건 전체를 파싱한 4계층 데이터 구조와 결정론적 검증은 기술적 기반이다. 기관형 수익모델과 정량 실증 목표도 마련했다. 현재 시제품은 충북 외국인 주민의 비자 전환과 정착 과정에 직접 적용되어 있다.",
    )

    # 11쪽 — 구현 상태
    add_page_break(document)
    document.add_heading("4. 제품(아이디어) 실현 방안", level=1)
    document.add_heading("4-1. 현재 구현 수준", level=2)
    add_caption(document, "그림 7. 기능별 구현 상태 점검표")
    add_standard_table(
        document,
        ["구분", "현재 구현 내용", "상태"],
        [
            ["데이터", "공식 원천문서 24건 전체 파싱, 원문·근거·규칙·서비스 4계층, 13개 관계형 테이블과 684개 근거 매핑", "구축 완료"],
            ["AI 챗봇", "의도·위험·사용자·지역·비자·언어 스크리닝, 제한형 데이터 조회, 다국어 답변", "시제품 적용"],
            ["위험 라우팅", "임금체불·산재·폭행·불법취업·거주조건 위반 분류, 사용자·지역별 기관 연결", "시제품 적용"],
            ["OCR 도우미", "서식 템플릿 분석, 필드별 누락·작성 주체·수동 확인, 6개 언어 질문 도우미", "시제품 적용"],
            ["공고 일정", "목표 비자별 접수기간 표시, 기준일·offset 계산, 일정 검색·개인 일정 등록", "시제품 적용"],
            ["안전·개인정보", "출처·연락처 원문 검증, OCR 비민감 필드 상태 활용, 대화 삭제와 비식별 운영 로그", "적용"],
        ],
        [1800, 6800, 1480],
        font_size=8.3,
        first_col_bold=True,
        center_columns={2},
    )
    document.add_heading("추진체계와 보유 역량", level=3)
    add_standard_table(
        document,
        ["역할", "수행 역량과 근거"],
        [
            ["행정 데이터", "공식 PDF·HWPX·HWP 24건 전체 파싱, 사람 검수, 4계층 구조·13개 관계형 테이블·684개 근거 매핑 운영"],
            ["웹·제품", "Next.js·Supabase 기반 반응형 화면, 다국어 온보딩, 문서·캘린더·기관 지도·OCR 구현"],
            ["품질·운영", "407개 자동 테스트, 트랜잭션 적재·롤백, CI·정적검사·PR 리뷰와 검증일 관리"],
        ],
        [2200, 7880],
        font_size=8.7,
        first_col_bold=True,
    )
    document.add_heading("기술 스택", level=3)
    add_body(
        document,
        "데이터 파이프라인은 Python 3.11+, pandas/polars, pdfplumber, pyhwp, lxml을 사용한다. 웹은 Next.js App Router·TypeScript·React·Tailwind CSS, 데이터베이스는 PostgreSQL 기반 Supabase, 배포는 Vercel을 사용한다. 품질관리는 pytest·Vitest·Ruff·ESLint·TypeScript typecheck·CI·PR 리뷰로 수행한다.",
        size=9.3,
        after=0,
    )

    # 12쪽 — 로드맵과 사업화
    add_page_break(document)
    document.add_heading("4-2. 단계별 개발 로드맵", level=2)
    add_standard_table(
        document,
        ["단계", "기간(안)", "주요 내용", "완료 기준"],
        [
            ["1단계: AI 행동연결", "2026.09~11", "챗봇 라우팅, OCR, 규칙·공고 일정 통합", "위험·서류·일정 경로별 검증"],
            ["2단계: 현장 실증", "2026.11~2027.02", "유학생·근로자·상담자 사용성 테스트", "2개 기관·사용자 30명 이상 실증"],
            ["3단계: 기관형 확장", "2027.03 이후", "운영자 검수, 통계, 대학·기업·지자체 도입", "유료 실증 1건 또는 도입의향·협약 3곳"],
        ],
        [2100, 1700, 3500, 2780],
        font_size=8.9,
        first_col_bold=True,
        center_columns={1},
    )
    document.add_heading("4-3. 사업화 모델", level=2)
    add_bullet(
        document,
        "개인 사용자(B2C): 요건 확인, 기본 체크리스트, 일정, 기관 검색을 무료로 제공해 접근 장벽을 낮춘다. 서류 묶음 관리·가족 공동 준비·고급 알림은 수요와 법률·개인정보 검토 후 선택형으로 확장한다.",
        size=9.4,
    )
    add_bullet(
        document,
        "기관·지자체(B2G/B2B): 대학 국제교류부서, 외국인지원센터, 기업 인사담당자, 지자체에 기관별 체크리스트 배포, 공고 변경 알림, 담당자 검수, 비식별 집계, 근거 링크를 제공한다. 기관용 연간 사용료로 AI 처리비와 공식 행정데이터의 검수·갱신·운영비를 충당한다.",
        size=9.4,
    )
    add_bullet(
        document,
        "확장 시장: 충북 지역특화형 비자와 유학생 체류·취업 안내에서 갱신 체계를 검증한 뒤, 타 지자체 지역특화형 비자와 전국 공통 체류자격으로 확대한다.",
        size=9.4,
    )
    add_standard_table(
        document,
        ["검증 항목", "1차 가설·목표", "검증 방법"],
        [
            ["초기 기관 접점", "채용박람회 공동주최 8개 대학·참여 32개 기업", "대학·기업·지원기관 문제 인터뷰 [3]"],
            ["가격 가설", "기관당 연 300만~1,200만원, 구축·연동 별도", "예산 보유 여부·지불의향·필수 기능 확인"],
            ["비용 효율", "OCR 1건 약 0.009달러·1,000건 약 9달러 예상 [13]", "실제 입력·출력 토큰과 완료 과업 수로 검증"],
            ["현장 실증", "2027년 2월까지 2개 기관·30명 이상", "과업 성공률·누락률·상담시간 전후 비교"],
            ["사업화 증거", "2027년 상반기 유료 실증 1건 또는 도입의향·협약 3곳", "계약·도입의향서·협약서로 확인"],
        ],
        [1900, 4570, 3610],
        font_size=8.2,
        first_col_bold=True,
    )
    add_body(
        document,
        "기관이 소액의 AI 처리비와 데이터 운영비를 부담하면 외국인 주민은 비용 없이 서류 누락·공고 일정·지원기관을 확인할 수 있다. 이를 통해 채용 이후 비자 전환 과정의 이탈을 줄이고, 충북 기업의 외국인 인력 유지와 지역 정착으로 편익을 환류한다. 가격과 비용은 실제 사용량과 현장 지불의향을 검증한 뒤 확정한다.",
    )

    # 13쪽 — 위험과 성과
    add_page_break(document)
    document.add_heading("4-4. 주요 위험과 해결 방안", level=2)
    add_standard_table(
        document,
        ["위험", "대응 방안"],
        [
            ["공고 변경·데이터 노후화", "적용기간·검증일·차수 관리, 변경 비교, 운영자 승인 후 배포"],
            ["원문 내부 충돌", "두 근거를 모두 보존하고 자동 판정을 차단해 담당기관 확인 상태로 표시"],
            ["AI 환각", "제한형 조회 결과만 사용하고 데이터에 없는 답변·연락처는 차단"],
            ["OCR 오인식", "원본 병기, 낮은 확신도 재촬영, 수동 확인, 최종 제출 전 사용자 확인"],
            ["개인정보 유출", "최소 수집, GPS 비저장, 민감정보 마스킹, 서버 키 분리, 보관기간 관리"],
            ["다국어 번역 오류", "핵심 용어집, 원문 병기, 전문 검수, 언어별 버전·검증일 관리"],
        ],
        [2600, 7480],
        font_size=9.0,
        first_col_bold=True,
    )
    document.add_heading("4-5. 성과 측정 계획", level=2)
    add_standard_table(
        document,
        ["지표", "측정 방법", "1차 실증 목표"],
        [
            ["위험 신호 미탐", "위험 표현·다국어 평가 세트", "0건"],
            ["기관 연락처 불일치", "데이터 원문과 문자열 비교", "0건"],
            ["공고 일정 출처 연결률", "캘린더 일정과 공고 근거 대조", "100%"],
            ["근거 없는 자동 일정", "공식일·기준일 없는 생성 검사", "0건"],
            ["필수서류 누락률", "OCR·체크리스트 모의 과업", "기존 대비 30% 감소"],
            ["정보 탐색시간", "공고문 직접 탐색군 비교", "40% 단축"],
            ["OCR 1건당 처리비", "실제 입력·출력 토큰과 API 단가", "예상치 약 0.009달러 검증"],
            ["AI 통합 과업 성공률", "질문→OCR→일정·기관 시나리오", "85% 이상"],
        ],
        [3700, 3900, 2480],
        font_size=8.2,
        first_col_bold=True,
        center_columns={2},
    )
    add_body(
        document,
        "위 수치는 현재 실적이 아니라 사업성과 검증을 위한 목표값이다. 현장 실증 전에 사용자군·표본 수·비교 조건을 확정한다.",
        size=9.5,
        italic=True,
        after=0,
    )

    # 14쪽 — 기대효과
    add_page_break(document)
    document.add_heading("5. 기대효과 및 활용분야", level=1)
    document.add_heading("5-1. 충북에 미칠 기대효과", level=2)
    add_bullet(document, "외국인 인재 정착: 취업·비자 전환·거주·행정 절차를 한 흐름으로 연결해 정보 부족에 따른 지역 이탈을 줄인다.")
    add_bullet(document, "기업 부담 완화: 고용주 요건과 제출서류를 신청자 서류와 분리하고 변경 기준을 제공해 중소기업의 행정 탐색 부담을 낮춘다.")
    add_bullet(document, "상담 품질 표준화: 대학·지원기관이 같은 공식 출처와 검증일을 공유하고 예외·충돌 사례를 검토 필요로 분리한다.")
    add_bullet(document, "데이터 기반 정책: 개인정보를 제외한 집계로 사용자가 막히는 단계와 반복 문의를 파악해 안내자료·상담 인력·지원사업을 개선한다.")
    add_bullet(document, "비용 대비 지역 환류: 기관이 부담하는 소액의 AI 처리비를 외국인 주민의 무료 비자 준비 지원으로 전환해 취업·비자 전환 이탈을 줄이고 기업 인력 유지와 충북 정착을 돕는다.")
    document.add_heading("5-2. 사회·경제적 파급효과", level=2)
    add_bullet(document, "행정 정보격차와 언어 장벽 완화")
    add_bullet(document, "신청 지연·재방문·서류 재발급 등 사회적 비용 절감")
    add_bullet(document, "외국인 인재의 안정적 취업과 지역 기업 인력난 완화 지원")
    add_bullet(document, "출처와 불확실성을 투명하게 보여주는 책임 있는 공공 AI 사례 제시")
    document.add_heading("5-3. 기타 활용분야", level=2)
    add_body(
        document,
        "타 지자체 지역특화형 비자 플랫폼, 대학 유학생 관리, 기업 외국인 인력 온보딩, 외국인지원기관 상담 보조, 행정문서 변경 모니터링, 쉬운 행정정보·다국어 콘텐츠로 확장할 수 있다.",
    )
    add_body(
        document,
        "비자부기는 외국인 주민의 질문, 서류, 공고 일정을 한곳에서 관리한다. AI가 질문의 맥락을 파악하고 OCR이 서류 상태를 확인한다. 위험 신호가 발견되면 전문기관 연락처를 우선 제시한다. 공식 문서를 근거·규칙·서비스 데이터로 연결한 충북형 비자·정착 지원 서비스다.",
    )

    # 15쪽 — 참고자료
    add_page_break(document)
    document.add_heading("참고자료 및 작성 근거", level=1)
    add_source(
        document,
        1,
        "체류외국인 주요통계",
        "법무부 출입국·외국인정책본부",
        "2025년 말 기준",
        "https://www.immigration.go.kr/moj/2412/subview.do",
    )
    add_source(
        document,
        2,
        "충북 지역특화형 비자 사업 숙련기능인력(E-7-4R) 모집 공고(2026년 6차)",
        "충청북도",
        "2026.06.09",
        "https://www.chungbuk.go.kr/www/selectGosiPblancView.do?key=422&no=68707",
    )
    add_source(
        document,
        3,
        "‘취업에서 정주까지’ 2026년 충청북도 외국인 유학생 채용박람회",
        "충청북도",
        "2026.04.28",
        "https://www.chungbuk.go.kr/www/selectBbsNttView.do?bbsNo=65&key=429&nttNo=420176",
    )
    add_source(
        document,
        4,
        "충북도, 도-시·군 외국인정책 협력회의 개최",
        "충청북도",
        "2026.04.08",
        "https://www.chungbuk.go.kr/www/selectBbsNttView.do?bbsNo=65&key=429&nttNo=418682",
    )
    add_source(
        document,
        5,
        "2025년 6월 출입국·외국인정책 통계월보",
        "법무부 출입국·외국인정책본부",
        "2025.06",
        "https://www.immigration.go.kr/bbs/immigration/227/483375/download.do",
    )
    add_source(
        document,
        6,
        "공고 원문 추출·검수·변경이력·데이터 적재 및 자동 테스트 결과",
        "비자부기 팀 visa-data 저장소",
        "2026.08.26 기준",
        "https://github.com/team-hansori/visa-data",
    )
    add_source(
        document,
        7,
        "온보딩·홈·캘린더·기관 지도·OCR·채팅 시제품 구현",
        "비자부기 팀 visa-bugi-web 저장소",
        "2026.08.28 기준",
        "https://github.com/team-hansori/visa-bugi-web",
    )
    add_source(
        document,
        8,
        "인재유치·지역발전·포용사회를 위한 2030 이민정책 미래전략",
        "법무부 출입국·외국인정책본부",
        "2026.03",
        "https://www.immigration.go.kr/bbs/immigration/214/491432/download.do",
    )
    add_source(
        document,
        9,
        "충북형 K-유학생 2만 명 유치 목표 관련 보도자료",
        "충청북도",
        "2026.04.29",
        "https://www.chungbuk.go.kr/www/selectBbsNttView.do?bbsNo=65&key=429&nttNo=420537",
    )
    add_source(
        document,
        10,
        "Study in Chungbuk 외국인 유학생 플랫폼",
        "충청북도",
        "2026년 확인",
        "https://www.studyinchungbuk.or.kr/",
    )
    add_source(
        document,
        11,
        "외국인을 위한 전자정부 하이코리아",
        "법무부 출입국·외국인정책본부",
        "2026년 확인",
        "https://www.hikorea.go.kr/",
    )
    add_source(
        document,
        12,
        "외국인 체류지 변경 신고 민원 안내",
        "정부24",
        "2026년 확인",
        "https://m.gov.kr/mw/AA020InfoCappView.do?CappBizCD=12700000026&HighCtgCD=A01010",
    )
    add_source(
        document,
        13,
        "GPT-5.4 Mini 모델 및 표준 API 요금",
        "OpenAI",
        "2026.08.28 확인",
        "https://developers.openai.com/api/docs/models/gpt-5.4-mini",
    )
    add_body(
        document,
        "통계와 정책 주장은 공식기관 자료를 기준으로 작성했다. 데이터·구현 수치는 팀 저장소 기준이며, 성과 목표는 현재 실적이 아닌 실증 목표다. 자격 결과는 참고용 자가진단으로 제공하고 최종 신청·판정은 관할 행정기관에서 수행한다.",
    )

    props = document.core_properties
    props.title = "제13회 전국 ICT융합 공모전 비자부기 사업계획서 최종안"
    props.subject = "디지털 시제품 · 충북 현안 해결 AI 혁신"
    props.author = "비자부기 팀"
    props.keywords = "비자부기, ICT융합 공모전, 사업계획서, 외국인, 지역특화형 비자"
    props.comments = "2026년 8월 28일 기준 제출용 최종안"
    document.save(output_path)


if __name__ == "__main__":
    output = Path(__file__).resolve().parents[1] / "2026_ICT융합공모전_비자부기_사업계획서_최종안.docx"
    build(output)
    print(output)
