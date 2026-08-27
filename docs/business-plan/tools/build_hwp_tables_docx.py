from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


FONT_NAME = "휴먼명조"
TABLE_WIDTH_DXA = 10080
TABLE_INDENT_DXA = 120
BLACK = "000000"
WHITE = "FFFFFF"
BRAND_GREEN = "173F36"


def set_run_font(run, size=11, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), FONT_NAME)


def format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.6):
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = line_spacing


def set_cell_text(cell, text, bold=False, size=11, color=BLACK, line_spacing=1.6):
    cell.text = ""
    lines = str(text).split("\n")
    paragraph = cell.paragraphs[0]
    format_paragraph(paragraph, line_spacing=line_spacing)
    for index, line in enumerate(lines):
        if index:
            run = paragraph.add_run()
            run.add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold)
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, color=WHITE):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = qn(f"w:{edge_name}")
        edge = tc_borders.find(tag)
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), edge_data.get("val", "single"))
        edge.set(qn("w:sz"), str(edge_data.get("sz", 8)))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), edge_data.get("color", BLACK))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def apply_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for tr in table._tbl.tr_lst:
        col_index = 0
        for tc in tr.tc_lst:
            tc_pr = tc.get_or_add_tcPr()
            grid_span = tc_pr.find(qn("w:gridSpan"))
            span = int(grid_span.get(qn("w:val"))) if grid_span is not None else 1
            cell_width = sum(widths[col_index : col_index + span])
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(cell_width))
            tc_w.set(qn("w:type"), "dxa")
            col_index += span


def style_table(table, header_rows=(0, 1), label_column=True):
    row_count = len(table.rows)
    col_count = len(table.columns)
    for r, row in enumerate(table.rows):
        prevent_row_split(row)
        for c, cell in enumerate(row.cells):
            set_cell_margins(cell)
            set_cell_fill(cell)
            outer_top = 14 if r == 0 else 8
            outer_bottom = 14 if r == row_count - 1 else 8
            outer_start = 14 if c == 0 else 8
            outer_end = 14 if c == col_count - 1 else 8
            set_cell_border(
                cell,
                top={"sz": outer_top},
                bottom={"sz": outer_bottom},
                start={"sz": outer_start},
                end={"sz": outer_end},
                insideH={"sz": 8},
                insideV={"sz": 8},
            )
            is_bold = r in header_rows or (label_column and c == 0)
            current_text = cell.text
            set_cell_text(cell, current_text, bold=is_bold)
    if len(table.rows) > 1:
        set_repeat_table_header(table.rows[1])


def add_caption(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(title)
    set_run_font(run, size=12, bold=True)


def add_structured_table(document, title, rows, widths, merges):
    add_caption(document, title)
    table = document.add_table(rows=len(rows), cols=len(widths))
    table.style = "Table Grid"
    for r, row_values in enumerate(rows):
        for c, value in enumerate(row_values):
            set_cell_text(table.cell(r, c), value)
    for r1, c1, r2, c2 in merges:
        merged = table.cell(r1, c1).merge(table.cell(r2, c2))
        set_cell_text(merged, rows[r1][c1], bold=r1 in (0, 1) or c1 == 0)
    apply_table_geometry(table, widths)
    style_table(table)
    return table


def add_compact_flow(document, title, steps, footer=None):
    add_caption(document, title)
    step_count = len(steps)
    if step_count == 5:
        box_width, arrow_width = 1776, 300
    elif step_count == 4:
        box_width, arrow_width = 2250, 360
    elif step_count == 3:
        box_width, arrow_width = 3060, 450
    else:
        raise ValueError("compact flow supports 3 to 5 steps")

    widths = []
    for index in range(step_count):
        widths.append(box_width)
        if index < step_count - 1:
            widths.append(arrow_width)
    assert sum(widths) == TABLE_WIDTH_DXA

    table = document.add_table(rows=2, cols=len(widths))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for step_index, (heading, body) in enumerate(steps):
        column = step_index * 2
        header_cell = table.cell(0, column)
        body_cell = table.cell(1, column)
        set_cell_fill(header_cell, BRAND_GREEN)
        set_cell_fill(body_cell, WHITE)
        set_cell_margins(header_cell, top=140, start=100, bottom=140, end=100)
        set_cell_margins(body_cell, top=180, start=100, bottom=180, end=100)
        set_cell_text(header_cell, heading, bold=True, color=WHITE, line_spacing=1.35)
        set_cell_text(body_cell, body, color=BLACK, line_spacing=1.45)
        border = {"val": "single", "sz": 10, "color": BLACK}
        set_cell_border(header_cell, top=border, bottom=border, start=border, end=border)
        set_cell_border(body_cell, top=border, bottom=border, start=border, end=border)

        if step_index < step_count - 1:
            arrow_column = column + 1
            arrow_cell = table.cell(0, arrow_column).merge(table.cell(1, arrow_column))
            set_cell_fill(arrow_cell, WHITE)
            set_cell_margins(arrow_cell, top=0, start=0, bottom=0, end=0)
            set_cell_text(arrow_cell, "→", bold=True, size=18, color=BRAND_GREEN, line_spacing=1.0)
            no_border = {"val": "nil", "sz": 0, "color": WHITE}
            set_cell_border(arrow_cell, top=no_border, bottom=no_border, start=no_border, end=no_border)

    prevent_row_split(table.rows[0])
    prevent_row_split(table.rows[1])
    apply_table_geometry(table, widths)

    if footer:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(7)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.3
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(footer)
        set_run_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor.from_string(BRAND_GREEN)
    return table


def add_page_break(document):
    document.add_page_break()


def configure_document(document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.6


def build(output_path):
    document = Document()
    configure_document(document)

    add_compact_flow(
        document,
        "그림 1. 비자부기 전체 서비스 흐름",
        [
            ("1. 사용자 확인", "언어·국적·지역\n관심 체류자격"),
            ("2. 상황 입력", "대화형 질문\n공문서·일정"),
            ("3. 규칙 처리", "AND/OR 계산\n검토 필요 분리"),
            ("4. 행동 안내", "서류·다음 할 일\n확정 일정 관리"),
            ("5. 기관 연결", "전화·길찾기\n공식 문의 경로"),
        ],
        "공식 근거와 검증일을 표시하고 사용자가 동의한 정보만 저장",
    )
    add_page_break(document)

    add_compact_flow(
        document,
        "그림 2. 공식 문서를 서비스 정보로 전환하는 데이터 신뢰 구조",
        [
            ("1. 공식 원문", "공고·지침·서식\n페이지·적용기간 보존"),
            ("2. 추출·사람 검수", "표·각주 대조\n충돌·미확인 기록"),
            ("3. 공통 데이터 구조", "13개 관계형 테이블\n조건·절차·서류·출처"),
            ("4. 서비스 제공", "규칙 계산·OCR\n다국어·기관 연결"),
        ],
        "생성형 AI가 자격을 추측하지 않고 검수된 규칙과 사람이 판단",
    )
    add_page_break(document)

    add_compact_flow(
        document,
        "그림 3. 비자 여정 단계별 서비스 책임 경계",
        [
            ("1. 추적", "요건·진행 현황\n비자부기 지원"),
            ("2. 준비", "서류·OCR·일정\n비자부기+사용자"),
            ("3. 제출", "공식 링크·제출처\n공식 접수기관"),
            ("4. 판정", "승인·불허·보완\n관할 행정기관"),
        ],
        "준비까지는 비자부기가 지원하고 제출·판정은 공식기관에서 수행",
    )
    add_page_break(document)

    add_compact_flow(
        document,
        "그림 4. 이주노동자의 E-7-4R 준비 과정",
        [
            ("1. 상황 파악", "체류기간·근무처\n연봉·한국어등급"),
            ("2. 요건 대조", "공식 기준 비교\n부족 항목 확인"),
            ("3. 일정 관리", "확정된 방문·제출\n일정만 관리"),
            ("4. 위험 확인", "계약조건 불일치\n일반 답변 중단"),
            ("5. 전문기관 연결", "노동·행정기관\n전화·위치 안내"),
        ],
        "자가진단·준비는 비자부기가 지원하고 신청·판정은 공식기관에서 수행",
    )
    add_page_break(document)

    add_compact_flow(
        document,
        "그림 5. 유학생의 유학→취업→정착 준비 과정",
        [
            ("1. 대상 확인", "학적·TOPIK\n허가 필요성"),
            ("2. 기준 안내", "허용시간\n제한업종"),
            ("3. 서류 이해", "OCR 사전 점검\n기한·금액·제출처"),
            ("4. 서류 준비", "체크리스트\n확정 일정 관리"),
            ("5. 장기 준비", "졸업 후 F-2-R\n요건 사전 점검"),
        ],
        "학업·취업 허가와 졸업 후 정착 준비를 하나의 연속된 여정으로 관리",
    )
    add_page_break(document)

    add_compact_flow(
        document,
        "그림 6. 현재 구현 상태와 다음 단계",
        [
            ("완료", "공통 데이터 구조\n온보딩·마이허브\nOCR 미리보기"),
            ("데모 구현", "진행 캘린더\n기관 지도\n반응형 흐름"),
            ("구현 예정", "규칙 계산 엔진\n다국어·근거 검색\n위험 상황 라우팅"),
        ],
        "1차 실증: 이주노동자 × E-7-4R 체류자격 트래커",
    )
    add_page_break(document)

    six_columns = [1300, 1756, 1756, 1756, 1756, 1756]
    five_columns = [1300, 2195, 2195, 2195, 2195]
    four_columns = [1300, 2927, 2927, 2926]

    add_structured_table(
        document,
        "상세 참고표 1. 비자부기 전체 서비스 흐름",
        [
            ["사용자 상황을 확인하고 근거 있는 다음 행동으로 연결", "", "", "", "", ""],
            ["구분", "1. 사용자 확인", "2. 상황 입력", "3. 규칙 처리", "4. 행동 안내", "5. 기관 연결"],
            ["입력", "언어·국적·지역\n관심 체류자격", "대화형 질문\n공문서 사진·일정", "구조화된 조건과\n사용자 입력 대조", "부족 요건·서류\n다음 할 일 제시", "지역·지원 분야에\n맞는 기관 선택"],
            ["처리 원칙", "로그인 전 선택값은\n현재 세션에만 보관\n최소 정보만 사용", "질문을 한 번에\n하나씩 제시하고\n불명확 항목 분리", "자격·점수는\nAND/OR 규칙으로\n결정론적으로 계산", "근거 없는 날짜와\n미확인 조건을\n임의로 추정하지 않음", "위험 신호에는\n일반 답변을 중단하고\n전문기관으로 연결"],
            ["결과", "개인화된 홈\n준비 현황", "확인된 사용자 상황\n검토 필요 항목", "충족·미충족\n검토 필요 구분", "서류 체크리스트\n확정 일정 관리", "전화·길찾기\n공식 문의 경로"],
            ["공통 원칙", "공식 근거·적용기간·마지막 검증일을 표시하고 사용자가 동의한 정보만 저장", "", "", "", ""],
        ],
        six_columns,
        [(0, 0, 0, 5), (5, 1, 5, 5)],
    )
    add_page_break(document)

    add_structured_table(
        document,
        "상세 참고표 2. 공식 문서를 서비스 정보로 전환하는 데이터 신뢰 구조",
        [
            ["원문 보존 → 항목 추출 → 사람 검수 → 구조화 → 서비스 제공", "", "", "", ""],
            ["구분", "1. 공식 원문", "2. 추출·검수", "3. 공통 데이터 구조", "4. 서비스 제공"],
            ["처리 내용", "법무부·충북도 공고\nPDF·HWPX·HWP\n원문·페이지 보존", "표·각주·붙임 분리\n사람이 원문 대조\n충돌·미확인 기록", "13개 관계형 테이블\nAND/OR 조건 논리\n절차·서류·출처 관리", "규칙 계산·OCR\n쉬운 설명·다국어\n체크리스트·기관 연결"],
            ["검증 항목", "공고 차수·적용기간\n문서 식별자\n페이지 기준", "표 병합셀·각주\n원문 내부 충돌\n미기재·미확인 구분", "UUID·FK·순환참조\n쿼터 산술·매핑 대상\n전체 트랜잭션 검증", "출처·검증일 표시\n검토 필요 분리\n최종 판정 대체 금지"],
            ["산출 결과", "검증 가능한\n원천 근거", "사람이 확인한\n항목별 데이터", "계산 가능한 규칙과\n변경 이력", "근거 있는 자가진단과\n다음 행동 안내"],
            ["핵심 원칙", "생성형 AI가 자격을 추측하지 않고 검수된 규칙과 사람이 판단", "", "", ""],
        ],
        five_columns,
        [(0, 0, 0, 4), (5, 1, 5, 4)],
    )
    add_page_break(document)

    add_structured_table(
        document,
        "상세 참고표 3. 비자 여정 단계별 서비스 책임 경계",
        [
            ["준비까지는 비자부기가 지원하고 제출·판정은 공식기관에서 수행", "", "", "", ""],
            ["구분", "1. 추적", "2. 준비", "3. 제출", "4. 판정"],
            ["주요 기능", "요건 충족도 확인\n규칙 기반 자가진단\n진행 현황 관리", "단계별 체크리스트\nOCR 사전 점검\n확정 일정 관리", "제출 절차 안내\n공식 링크·제출처\n실제 접수에는 미관여", "공식 승인·불허\n보완 요구·결과 통지\nAI 판단 금지"],
            ["처리 주체", "비자부기\n+ 사용자", "비자부기\n+ 사용자·고용주", "사용자·고용주\n+ 공식 접수기관", "법무부 출입국·\n외국인관서 등"],
            ["책임 범위", "근거 있는 사전 점검\n공식 결정 대체 아님", "준비 과정 지원\n최종 제출 전 사용자 확인", "공식 채널에서 접수\n비자부기 책임 범위 밖", "법령·행정 기준에 따른\n최종 판단"],
        ],
        five_columns,
        [(0, 0, 0, 4)],
    )
    add_page_break(document)

    add_structured_table(
        document,
        "상세 참고표 4. 이주노동자의 E-7-4R 준비 과정",
        [
            ["응우옌 반 A · 27세 · 베트남 · E-9 · 음성군 제조업체 근무 2년 차", "", "", "", "", ""],
            ["구분", "1. 상황 파악", "2. 요건 대조", "3. 일정 관리", "4. 위험 확인", "5. 전문기관 연결"],
            ["사용자 행동", "체류기간·근무처\n연봉·한국어등급 입력", "공식 E-7-4R 기준과\n현재 조건 비교", "동의한 경우에만\n방문·제출 일정 추가", "계약조건 불일치 등\n위험 신호 확인", "관할 전문기관의\n연락처·위치 확인"],
            ["비자부기 지원", "입력할 조건을\n질문 순서로 안내\n검토 필요 항목 분리", "AND/OR 규칙으로\n충족·미충족·검토\n필요 결과를 구분", "근거 없는 상대일정을\n추정하지 않고\n확정 일정만 관리", "일반 답변을 중단하고\nAI의 임의 판단 없이\n위험 유형을 분류", "노동·행정·긴급지원\n기관으로 즉시 연결"],
            ["판정·제출 주체", "사용자 입력\n필요 시 고용주 확인", "규칙 기반 사전 점검\n공식 결정은 아님", "사용자 선택·관리\n좌표·일정 최소 수집", "AI 자동 판정 금지\n전문기관 상담 권고", "관할 노동·행정기관의\n최종 안내·처리"],
            ["책임 경계", "자가진단·준비는 비자부기에서 지원하고 신청·판정은 공식기관에서 수행", "", "", "", ""],
        ],
        six_columns,
        [(0, 0, 0, 5), (5, 1, 5, 5)],
    )
    add_page_break(document)

    add_structured_table(
        document,
        "상세 참고표 5. 유학생의 유학→취업→정착 준비 과정",
        [
            ["바트 체첵 · 22세 · 몽골 · D-2 · 충북 소재 대학 재학", "", "", "", "", ""],
            ["구분", "1. 대상 확인", "2. 기준 안내", "3. 서류 이해", "4. 서류 준비", "5. 장기 준비"],
            ["사용자 행동", "학적·TOPIK·허가\n필요성 확인", "허용시간·제한업종\n기준 확인", "고지서 항목과\n기한·금액 확인", "신청서 작성 후\n학교 제출", "졸업 후 F-2-R\n요건 사전 점검"],
            ["비자부기 지원", "질문형 온보딩으로\n현재 상황 구분\n확인할 기준 제시", "공식 기준을 쉬운\n한국어로 설명\n출처·검증일 표시", "OCR 결과를 항목별\n확인 필요·누락으로\n구분해 재검토 지원", "서류 체크리스트와\n확정 일정만\n캘린더에 관리", "취업·정착 준비를\n후속 여정으로\n연결해 안내"],
            ["판정·제출 주체", "사용자 확인\n필요 시 학교 문의", "비자부기 안내\n공식 기준 병기", "사용자 최종 확인\n자동 판정 금지", "대학 국제교류부서\n또는 출입국기관", "관할 행정기관의\n최종 판단"],
            ["핵심 가치", "학업·취업 허가와 졸업 후 정착 준비를 하나의 연속된 여정으로 관리", "", "", "", ""],
        ],
        six_columns,
        [(0, 0, 0, 5), (5, 1, 5, 5)],
    )
    add_page_break(document)

    add_structured_table(
        document,
        "상세 참고표 6. 현재 구현 상태와 다음 단계",
        [
            ["완료·데모·구현 예정 항목을 구분해 과장 없이 표시", "", "", ""],
            ["구분", "완료", "데모 구현", "구현 예정"],
            ["주요 범위", "공통 데이터 구조 13개 테이블\nAND/OR 자격조건 논리\n온보딩·마이허브 화면\nOCR 이미지 선택·미리보기", "비자 진행 캘린더\n지도 기반 기관 안내\n반응형 사용자 흐름\n전화·길찾기 화면", "규칙 기반 계산 엔진\n다국어 설명·근거 검색\n위험 상황 기관 라우팅\nOCR 분석·민감정보 마스킹"],
            ["현재 수준", "설계·검수 확정\n웹 화면에 골격 구현", "사용 흐름 확인 가능\n일부 실데이터 연동 전", "현장 검증 후 단계적 구현\n공식 판정 대체 기능 없음"],
            ["1차 실증", "이주노동자 × E-7-4R 체류자격 트래커를 우선 검증", "", ""],
        ],
        four_columns,
        [(0, 0, 0, 3), (4, 1, 4, 3)],
    )

    core_props = document.core_properties
    core_props.title = "2026 ICT융합공모전 비자부기 한글 복붙용 도식·표 모음"
    core_props.subject = "사업계획서 간략 도식 및 상세표 자료"
    core_props.author = "비자부기 팀"
    core_props.keywords = "비자부기, 사업계획서, 도식, 표"
    document.save(output_path)


if __name__ == "__main__":
    output = Path(__file__).resolve().parents[1] / "2026_ICT융합공모전_한글복붙용_표모음.docx"
    build(output)
    print(output)
